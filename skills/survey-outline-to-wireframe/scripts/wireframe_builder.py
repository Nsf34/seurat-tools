"""
Survey wireframe builder for the current Merck client-facing format.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


TITLE_BLUE = RGBColor(0x1B, 0x2E, 0x59)
SUBTITLE_ORANGE = RGBColor(0xF1, 0x5B, 0x2A)
DATE_GRAY = RGBColor(0x75, 0x8B, 0x97)
SECTION_BLUE = "002060"
SECTION_BLUE_ALT = "1B2E59"
COLUMN_BLUE = "4970C8"
H2_BLUE = RGBColor(0x28, 0xA8, 0xE0)
RED = RGBColor(0xEE, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

TITLE_FONT = "Franklin Gothic Demi"
BODY_FONT = "Franklin Gothic Book"
TITLE_SIZE = Pt(14)
SUBTITLE_SIZE = Pt(14)
DATE_SIZE = Pt(12)
HEADING1_SIZE = Pt(18)
HEADING2_SIZE = Pt(14)
BODY_SIZE = Pt(11)
PROG_NOTE_SIZE = Pt(10)

COL_WIDTHS = [1379, 4871, 2662]


class WireframeBuilder:
    def __init__(self):
        self.doc = Document()
        self._current_table = None
        self._section_count = 0
        self._question_block_index = 0
        self._current_question_mode = None
        self._current_question_counter = 0
        self._setup_document()

    def _setup_document(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = BODY_SIZE
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

        for style_name in [
            "List Paragraph",
            "List Bullet",
            "List Bullet 2",
            "List Bullet 3",
            "List Number",
            "List Number 2",
            "List Number 3",
        ]:
            if style_name in self.doc.styles:
                list_style = self.doc.styles[style_name]
                list_style.font.name = BODY_FONT
                list_style.font.size = BODY_SIZE
                list_style.paragraph_format.space_before = Pt(0)
                list_style.paragraph_format.space_after = Pt(0)

        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _clear_paragraph(self, paragraph):
        element = paragraph._element
        for child in list(element):
            if child.tag != qn("w:pPr"):
                element.remove(child)

    def _add_run(
        self,
        paragraph,
        text,
        bold=False,
        italic=False,
        color=None,
        size=None,
        font_name=None,
    ):
        run = paragraph.add_run(text)
        run.font.name = font_name or BODY_FONT
        run.font.size = size or BODY_SIZE
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return run

    def _add_parts(self, paragraph, parts, default_color=None, default_size=None):
        for part in parts:
            if isinstance(part, str):
                self._add_run(paragraph, part, color=default_color, size=default_size)
                continue

            text = part[0]
            fmt = part[1] if len(part) > 1 else False

            bold = False
            italic = False
            color = default_color

            if fmt in (False, None, "plain"):
                pass
            elif fmt in (True, "var"):
                if color is None:
                    color = RED
            elif fmt == "bold":
                bold = True
            elif fmt == "italic":
                italic = True
            elif fmt == "italic_var":
                italic = True
                if color is None:
                    color = RED
            elif fmt == "bold_var":
                bold = True
                if color is None:
                    color = RED
            elif fmt == "bold_italic":
                bold = True
                italic = True
            elif fmt == "bold_italic_var":
                bold = True
                italic = True
                if color is None:
                    color = RED

            self._add_run(
                paragraph,
                text,
                bold=bold,
                italic=italic,
                color=color,
                size=default_size,
            )

    def _add_text_or_parts(self, paragraph, value, default_color=None, default_size=None):
        if isinstance(value, str):
            self._add_run(paragraph, value, color=default_color, size=default_size)
            return
        if isinstance(value, list):
            self._add_parts(paragraph, value, default_color=default_color, default_size=default_size)
            return
        self._add_run(paragraph, str(value), color=default_color, size=default_size)

    def _set_cell_shading(self, cell, fill_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        for existing in tc_pr.findall(qn("w:shd")):
            tc_pr.remove(existing)
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>'
        )
        tc_pr.append(shading)

    def _set_cell_width(self, cell, width_dxa):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
            tc_pr.append(tc_w)
        else:
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")

    def _set_fixed_table_layout(self, table):
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
            tbl_pr.append(layout)
        else:
            layout.set(qn("w:type"), "fixed")

    def _set_table_grid(self, table):
        tbl = table._tbl
        tbl_grid = tbl.find(qn("w:tblGrid"))
        if tbl_grid is None:
            tbl_grid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
            tbl.insert(1, tbl_grid)

        for existing in list(tbl_grid):
            tbl_grid.remove(existing)

        for width in COL_WIDTHS:
            grid_col = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{width}"/>')
            tbl_grid.append(grid_col)

    def _prepare_table(self, table):
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.autofit = False
        self._set_fixed_table_layout(table)
        self._set_table_grid(table)

    def _set_row_widths(self, row):
        for idx, width in enumerate(COL_WIDTHS):
            self._set_cell_width(row.cells[idx], width)

    def _add_centered_multiline_paragraph(self, text, color, size, font_name, bold=False):
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(6)

        lines = str(text).splitlines() or [""]
        for index, line in enumerate(lines):
            run = self._add_run(
                paragraph,
                line,
                bold=bold,
                color=color,
                size=size,
                font_name=font_name,
            )
            if index < len(lines) - 1:
                run.add_break(WD_BREAK.LINE)
        return paragraph

    def _add_overview_heading(self, text):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        self._add_run(
            paragraph,
            text,
            bold=True,
            color=TITLE_BLUE,
            size=HEADING1_SIZE,
            font_name=TITLE_FONT,
        )
        return paragraph

    def _add_overview_subheading(self, text):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        self._add_run(
            paragraph,
            text,
            bold=True,
            color=H2_BLUE,
            size=HEADING2_SIZE,
            font_name=TITLE_FONT,
        )
        return paragraph

    def _get_list_style_name(self, kind="bullet", level=0):
        if kind == "number":
            candidates = ["List Number", "List Number 2", "List Number 3"]
        else:
            candidates = ["List Bullet", "List Bullet 2", "List Bullet 3"]

        index = min(max(level, 0), len(candidates) - 1)
        style_name = candidates[index]
        if style_name in self.doc.styles:
            return style_name
        return "List Paragraph"

    def _new_list_paragraph(self, container=None, kind="bullet", level=0):
        container = container or self.doc
        paragraph = container.add_paragraph()
        style_name = self._get_list_style_name(kind=kind, level=level)
        try:
            paragraph.style = style_name
        except KeyError:
            paragraph.style = "List Paragraph"
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return paragraph

    def _add_list_paragraph(self, value, bold_prefix=None, container=None, kind="bullet", level=0):
        paragraph = self._new_list_paragraph(container=container, kind=kind, level=level)
        if bold_prefix:
            self._add_run(paragraph, bold_prefix, bold=True)
            if value:
                spacer = "" if str(value).startswith(" ") else " "
                self._add_run(paragraph, f"{spacer}{value}")
        else:
            self._add_text_or_parts(paragraph, value)
        return paragraph

    def _add_hierarchical_quota_items(self, quotas):
        current_group_open = False

        for quota in quotas:
            if isinstance(quota, tuple) and len(quota) == 2:
                label, value = quota
                if str(value).strip():
                    self._add_list_paragraph(value, bold_prefix=label, kind="bullet", level=0)
                    current_group_open = False
                else:
                    self._add_list_paragraph([(label, "bold")], kind="bullet", level=0)
                    current_group_open = True
                continue

            if isinstance(quota, dict):
                label = quota.get("label")
                value = quota.get("value", "")
                children = quota.get("children") or quota.get("bullets") or quota.get("items") or []

                if label and str(value).strip():
                    self._add_list_paragraph(value, bold_prefix=label, kind="bullet", level=0)
                elif label:
                    self._add_list_paragraph([(label, "bold")], kind="bullet", level=0)
                elif value:
                    self._add_list_paragraph(value, kind="bullet", level=0)

                for child in children:
                    if isinstance(child, tuple) and len(child) == 2:
                        self._add_list_paragraph(child[1], bold_prefix=child[0], kind="bullet", level=1)
                    else:
                        self._add_list_paragraph(child, kind="bullet", level=1)

                current_group_open = bool(label and not str(value).strip())
                continue

            self._add_list_paragraph(
                quota,
                kind="bullet",
                level=1 if current_group_open else 0,
            )

    def _begin_question_block(self, section_name):
        self._current_question_counter = 0
        lower_name = section_name.lower()

        if "screener" in lower_name:
            self._current_question_mode = ("prefix", "S")
            return

        if "demographic" in lower_name or "demographics" in lower_name or "profiling" in lower_name:
            self._current_question_mode = ("prefix", "D")
            return

        self._question_block_index += 1
        self._current_question_mode = ("block", self._question_block_index)

    def _next_question_number(self):
        if self._current_question_mode is None:
            self._begin_question_block("Section")

        self._current_question_counter += 1
        mode, value = self._current_question_mode

        if mode == "prefix":
            return f"{value}{self._current_question_counter}."

        return f"Q{value}{self._current_question_counter:02d}."

    def set_cover(self, title, subtitle="Survey Wireframe", date=""):
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        title_paragraph = self._add_centered_multiline_paragraph(
            title,
            color=TITLE_BLUE,
            size=TITLE_SIZE,
            font_name=TITLE_FONT,
            bold=False,
        )
        title_paragraph.paragraph_format.space_before = Pt(48)
        title_paragraph.paragraph_format.space_after = Pt(12)

        subtitle_paragraph = self._add_centered_multiline_paragraph(
            subtitle,
            color=SUBTITLE_ORANGE,
            size=SUBTITLE_SIZE,
            font_name=TITLE_FONT,
            bold=False,
        )
        subtitle_paragraph.paragraph_format.space_after = Pt(12)

        if date:
            date_paragraph = self._add_centered_multiline_paragraph(
                date,
                color=DATE_GRAY,
                size=DATE_SIZE,
                font_name=BODY_FONT,
                bold=False,
            )
            date_paragraph.paragraph_format.space_after = Pt(24)

    def add_overview(
        self,
        objectives=None,
        quotas=None,
        survey_flow=None,
        survey_flow_description=None,
    ):
        self._add_overview_heading("Survey Overview")

        if objectives:
            self._add_overview_subheading("Quant Research Objectives")
            for objective in objectives:
                self._add_list_paragraph(objective, kind="bullet", level=0)

        if quotas:
            self._add_overview_subheading("Quotas")
            self._add_hierarchical_quota_items(quotas)

        if survey_flow:
            self._add_overview_subheading("Survey Flow")
            if survey_flow_description:
                paragraph = self.doc.add_paragraph()
                self._add_run(paragraph, survey_flow_description)

            for flow_item in survey_flow:
                if isinstance(flow_item, dict):
                    section = flow_item.get("section", "")
                    description = flow_item.get("description", "")
                    bullets = flow_item.get("bullets", [])

                    paragraph = self._new_list_paragraph(kind="number", level=0)
                    self._add_run(paragraph, section, bold=True)
                    if description:
                        detail = self._new_list_paragraph(kind="bullet", level=1)
                        self._add_text_or_parts(detail, description)

                    for bullet in bullets:
                        bullet_paragraph = self._new_list_paragraph(kind="bullet", level=1)
                        self._add_text_or_parts(bullet_paragraph, bullet)
                elif isinstance(flow_item, tuple) and len(flow_item) == 2:
                    heading = self._new_list_paragraph(kind="number", level=0)
                    self._add_run(heading, flow_item[0], bold=True)
                    self._add_list_paragraph(flow_item[1], kind="bullet", level=1)
                else:
                    self._add_list_paragraph(flow_item, kind="number", level=0)

        self.doc.add_paragraph("")

    def _add_section_header_row(self, section_name, header_fill):
        row = self._current_table.add_row()
        self._set_row_widths(row)
        merged = row.cells[0].merge(row.cells[2])
        self._set_cell_shading(merged, header_fill)
        paragraph = merged.paragraphs[0]
        self._clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(
            paragraph,
            section_name,
            bold=True,
            color=WHITE,
            font_name=TITLE_FONT,
        )

    def _add_column_header_row(self):
        row = self._current_table.add_row()
        self._set_row_widths(row)
        headers = ["Q #", "Survey Question", "Question Objective"]

        for idx, header in enumerate(headers):
            cell = row.cells[idx]
            self._set_cell_shading(cell, COLUMN_BLUE)
            paragraph = cell.paragraphs[0]
            self._clear_paragraph(paragraph)
            self._add_run(
                paragraph,
                header,
                bold=True,
                color=WHITE,
                font_name=TITLE_FONT,
            )

    def _add_section_objectives_row(self, objectives):
        row = self._current_table.add_row()
        self._set_row_widths(row)
        merged = row.cells[0].merge(row.cells[2])
        paragraph = merged.paragraphs[0]
        self._clear_paragraph(paragraph)
        self._add_run(paragraph, "Section Objectives")

        for objective in objectives:
            self._add_list_paragraph(objective, container=merged, kind="bullet", level=0)

    def start_section(self, section_name, objectives=None, header_fill=None):
        if self._section_count > 0:
            self.doc.add_paragraph("")

        table = self.doc.add_table(rows=0, cols=3)
        self._prepare_table(table)
        self._current_table = table
        self._section_count += 1

        self._begin_question_block(section_name)
        self._add_section_header_row(section_name, header_fill or SECTION_BLUE)
        self._add_column_header_row()

        if objectives:
            self._add_section_objectives_row(objectives)

        return table

    def add_subsection_header(self, section_name, objectives=None, header_fill=None):
        if self._current_table is None:
            raise ValueError("Must call start_section() before add_subsection_header().")

        self._begin_question_block(section_name)
        self._add_section_header_row(section_name, header_fill or SECTION_BLUE)
        self._add_column_header_row()

        if objectives:
            self._add_section_objectives_row(objectives)

    def add_subsection(self, section_name, objectives=None, header_fill=None):
        self.add_subsection_header(
            section_name,
            objectives=objectives,
            header_fill=header_fill,
        )

    def add_message(self, text, message_label="Message"):
        if self._current_table is None:
            raise ValueError("Must call start_section() before add_message().")

        row = self._current_table.add_row()
        self._set_row_widths(row)

        label_paragraph = row.cells[0].paragraphs[0]
        self._clear_paragraph(label_paragraph)
        self._add_run(label_paragraph, message_label)

        merged = row.cells[1].merge(row.cells[2])
        paragraph = merged.paragraphs[0]
        self._clear_paragraph(paragraph)
        self._add_text_or_parts(paragraph, text)

    def add_question(
        self,
        title,
        question_parts,
        instruction=None,
        response_lines=None,
        notes=None,
        objective="",
        programming_notes=None,
        sub_questions=None,
        question_number="",
    ):
        if self._current_table is None:
            raise ValueError("Must call start_section() before add_question().")

        row = self._current_table.add_row()
        self._set_row_widths(row)

        qnum_paragraph = row.cells[0].paragraphs[0]
        self._clear_paragraph(qnum_paragraph)
        qnum_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if question_number is False:
            resolved_question_number = ""
        elif question_number in (None, "", "auto"):
            resolved_question_number = self._next_question_number()
        else:
            resolved_question_number = str(question_number)

        if resolved_question_number:
            self._add_run(qnum_paragraph, resolved_question_number)

        question_cell = row.cells[1]
        question_paragraph = question_cell.paragraphs[0]
        self._clear_paragraph(question_paragraph)
        self._add_run(question_paragraph, f"{title}: ", bold=True)
        self._add_text_or_parts(question_paragraph, question_parts)
        if instruction:
            if question_paragraph.text and not question_paragraph.text.endswith(" "):
                self._add_run(question_paragraph, " ")
            self._add_run(question_paragraph, instruction, italic=True)

        for sub_question in sub_questions or []:
            if sub_question.get("separator", True):
                question_cell.add_paragraph("")
            sub_paragraph = question_cell.add_paragraph()
            self._add_text_or_parts(sub_paragraph, sub_question.get("parts", []))
            if sub_question.get("instruction"):
                if sub_paragraph.text and not sub_paragraph.text.endswith(" "):
                    self._add_run(sub_paragraph, " ")
                self._add_run(sub_paragraph, sub_question["instruction"], italic=True)

        for response in response_lines or []:
            response_paragraph = self._new_list_paragraph(container=question_cell, kind="bullet", level=0)
            if isinstance(response, str):
                self._add_run(response_paragraph, response)
            elif isinstance(response, dict):
                if "bold_prefix" in response:
                    self._add_run(response_paragraph, response["bold_prefix"], bold=True)
                    text = response.get("text", "")
                    if text:
                        spacer = "" if text.startswith(" ") else " "
                        self._add_run(response_paragraph, f"{spacer}{text}")
                elif "parts" in response:
                    self._add_parts(response_paragraph, response["parts"])
                else:
                    self._add_run(response_paragraph, response.get("text", ""))

                if response.get("italic"):
                    for run in response_paragraph.runs:
                        run.italic = True
                if response.get("italic_note"):
                    self._add_run(
                        response_paragraph,
                        f" ({response['italic_note']})",
                        italic=True,
                    )
                if response.get("var_tag"):
                    self._add_run(response_paragraph, f" {response['var_tag']}", color=RED)
            else:
                self._add_run(response_paragraph, str(response))

        for note in notes or []:
            note_paragraph = self._new_list_paragraph(container=question_cell, kind="bullet", level=0)
            if isinstance(note, str):
                self._add_run(note_paragraph, note, italic=True)
            elif isinstance(note, dict) and "parts" in note:
                self._add_parts(note_paragraph, note["parts"])
            else:
                self._add_run(note_paragraph, str(note), italic=True)

        objective_cell = row.cells[2]
        objective_paragraph = objective_cell.paragraphs[0]
        self._clear_paragraph(objective_paragraph)
        if objective:
            self._add_text_or_parts(objective_paragraph, objective)

        if objective and programming_notes:
            objective_cell.add_paragraph("")

        for note in programming_notes or []:
            style = "Normal"
            if isinstance(note, dict) and note.get("is_bullet"):
                style = "bullet"
            if style == "bullet":
                note_paragraph = self._new_list_paragraph(
                    container=objective_cell,
                    kind="bullet",
                    level=note.get("level", 0) if isinstance(note, dict) else 0,
                )
            else:
                note_paragraph = objective_cell.add_paragraph(style="Normal")

            if isinstance(note, str):
                self._add_run(note_paragraph, note, color=RED, size=PROG_NOTE_SIZE)
            elif isinstance(note, dict):
                if "parts" in note:
                    self._add_parts(
                        note_paragraph,
                        note["parts"],
                        default_color=RED,
                        default_size=PROG_NOTE_SIZE,
                    )
                else:
                    self._add_run(
                        note_paragraph,
                        note.get("text", ""),
                        bold=note.get("bold", False),
                        italic=note.get("italic", False),
                        color=RED,
                        size=PROG_NOTE_SIZE,
                    )
            else:
                self._add_run(note_paragraph, str(note), color=RED, size=PROG_NOTE_SIZE)

    def add_qualification_message(self, text):
        self.add_message(text, message_label="Qualification Message")

    def add_termination_message(self, text):
        self.add_message(text, message_label="Termination Message")

    def add_freeform_note(self, text):
        paragraph = self.doc.add_paragraph()
        self._add_text_or_parts(paragraph, text)

    def save(self, filepath):
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        self.doc.save(filepath)
        print(f"Wireframe saved: {filepath}")
        return filepath


if __name__ == "__main__":
    builder = WireframeBuilder()
    builder.set_cover(
        title="Test Study\nPath to Purchase & PDH",
        subtitle="Survey Wireframe",
        date="March 2026",
    )
    builder.add_overview(
        objectives=[
            "Map the shopper journey.",
            "Identify key decision drivers.",
        ],
        quotas=[
            "N = 1,000",
            ("Markets:", "US, UK"),
        ],
        survey_flow=[
            {"section": "Screener", "bullets": ["Qualification and assignment"]},
            {"section": "Plan", "bullets": ["Planning behavior"]},
        ],
        survey_flow_description="This survey includes two sections:",
    )
    builder.start_section(
        "Screener",
        objectives=[
            "Ensure the right respondents qualify.",
            "Assign key variables.",
        ],
    )
    builder.add_message("Thank you for taking part in this survey.")
    builder.add_question(
        title="Age",
        question_parts=[("What is your age?", False)],
        instruction="Select one.",
        response_lines=["Dropdown with range from 18 - Over 80"],
        objective="Ensure respondent meets age qualifications.",
        programming_notes=["Terminate under 18."],
    )
    builder.save("test_wireframe.docx")
