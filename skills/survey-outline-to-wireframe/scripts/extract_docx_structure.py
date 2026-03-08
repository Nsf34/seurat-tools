"""
Extract a readable structural summary from a .docx file.

Usage:
    python extract_docx_structure.py input.docx
    python extract_docx_structure.py input.docx -o output.txt
"""

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def paragraph_summary(paragraph):
    text = paragraph.text.replace("\n", " // ").strip()
    style = paragraph.style.name if paragraph.style else ""
    return f"PARA [{style}]: {text}"


def cell_text(cell):
    parts = []
    for paragraph in cell.paragraphs:
        text = paragraph.text.replace("\n", " // ").strip()
        if text:
            parts.append(text)
    return " | ".join(parts)


def cell_span(cell):
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return 1
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return 1
    return int(grid_span.get(qn("w:val"), "1"))


def table_summary(table, index):
    lines = []
    style = table.style.name if table.style else ""
    lines.append(f"TABLE {index}: rows={len(table.rows)} cols={len(table.columns)} style={style}")
    for row_index, row in enumerate(table.rows):
        cell_parts = []
        for cell_index, cell in enumerate(row.cells):
            text = cell_text(cell)
            span = cell_span(cell)
            span_text = f" [span={span}]" if span > 1 else ""
            cell_parts.append(f"C{cell_index}: {text}{span_text}")
        lines.append(f"ROW {row_index}: " + " || ".join(cell_parts))
    return lines


def build_summary(docx_path):
    doc = Document(docx_path)
    lines = []
    lines.append(f"FILE: {docx_path}")
    lines.append(f"PARAGRAPHS: {len(doc.paragraphs)}")
    lines.append(f"TABLES: {len(doc.tables)}")
    lines.append("")
    lines.append("BODY")
    lines.append("-" * 80)

    for paragraph in doc.paragraphs:
        lines.append(paragraph_summary(paragraph))

    if doc.tables:
        lines.append("")
        lines.append("TABLES")
        lines.append("-" * 80)
        for index, table in enumerate(doc.tables, start=1):
            lines.extend(table_summary(table, index))
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path")
    parser.add_argument("-o", "--output")
    args = parser.parse_args()

    source = Path(args.docx_path)
    summary = build_summary(source)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = source.with_suffix(".txt")

    output_path.write_text(summary, encoding="utf-8")
    print(output_path)


if __name__ == "__main__":
    main()
