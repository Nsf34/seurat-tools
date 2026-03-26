"""
Survey Document Builder — python-docx engine for Seurat Group survey documents.

Generates a correctly formatted, programmer-ready survey document (.docx) from
structured question data. Used by the survey-wireframe-to-doc skill.

Key formatting rules:
  - Font: Franklin Gothic Book, 11pt
  - Header + response rows in ONE merged table (R0-R3 header spans all cols)
  - Standalone 4x1 header tables only for messages, dropdowns, and open-ends
  - RED (FF0000): topic labels, question numbers, programming notes, piped
    variables in question text, column labels (C1:, C2:), logic notes, NEW SCREEN
  - BLACK: question text, selection instructions (italic), response option text,
    column header descriptions, scale labels
  - Tables: BFBFBF borders, blue (B4C6E7) shading on R0 (topic row)
  - Empty coding cells pre-formatted as red so manual edits auto-appear red

Usage:
    from survey_doc_builder import SurveyDocBuilder

    builder = SurveyDocBuilder()
    builder.set_study_info(title="Study Name", date="March 2026")
    builder.add_overview(objectives=[...], structure=[...], criteria=[...])
    builder.add_section_header("Screener", objectives=[...])
    builder.add_message("Introduction Message", "M1.", "Thank you...", "Show to all.")
    builder.add_question(q_number="S1.", topic="Age", ...)
    builder.save("output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import re

# ── Constants ────────────────────────────────────────────────

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
BLUE_FILL = "B4C6E7"       # Light blue for header R0 shading
BORDER_COLOR = "BFBFBF"    # Light gray for all table borders
FONT_NAME = "Franklin Gothic Book"
FONT_SIZE = Pt(11)         # 139700 EMU
HEADING2_COLOR = RGBColor(0x2F, 0x54, 0x96)

# Column widths (dxa / twentieths of a point)
COL0_WIDTH = 576     # Narrow blank/number column
COL1_WIDTH = 6031    # Wide option text column
COL2_WIDTH = 2748    # Medium notes column


class SurveyDocBuilder:
    """Builds a formatted survey document .docx matching Seurat standards."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_numbering()
        self._first_block = True

    # ── Setup ────────────────────────────────────────────────

    def _setup_styles(self):
        """Configure document-level styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE

        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        for section in self.doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

    def _setup_numbering(self):
        """Create a multi-level bullet numbering definition for proper Word lists."""
        try:
            numbering_part = self.doc.part.numbering_part
            numbering_elm = numbering_part.element
        except Exception:
            self._bullet_num_id = 1
            return

        # Find next available IDs
        abstract_nums = numbering_elm.findall(qn('w:abstractNum'))
        next_abstract = max(
            (int(an.get(qn('w:abstractNumId'), '0')) for an in abstract_nums),
            default=-1
        ) + 1

        nums = numbering_elm.findall(qn('w:num'))
        next_num = max(
            (int(n.get(qn('w:numId'), '0')) for n in nums),
            default=0
        ) + 1

        # 3-level bullet definition (circle, open circle, small square)
        abstract_xml = (
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{next_abstract}">'
            f'<w:multiLevelType w:val="hybridMultilevel"/>'
            f'<w:lvl w:ilvl="0"><w:start w:val="1"/>'
            f'<w:numFmt w:val="bullet"/>'
            f'<w:lvlText w:val="\u2022"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
            f'</w:lvl>'
            f'<w:lvl w:ilvl="1"><w:start w:val="1"/>'
            f'<w:numFmt w:val="bullet"/>'
            f'<w:lvlText w:val="o"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr>'
            f'<w:rPr><w:rFonts w:ascii="Courier New" w:hAnsi="Courier New"/></w:rPr>'
            f'</w:lvl>'
            f'<w:lvl w:ilvl="2"><w:start w:val="1"/>'
            f'<w:numFmt w:val="bullet"/>'
            f'<w:lvlText w:val="\u25AA"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr>'
            f'</w:lvl>'
            f'</w:abstractNum>'
        )

        num_xml = (
            f'<w:num {nsdecls("w")} w:numId="{next_num}">'
            f'<w:abstractNumId w:val="{next_abstract}"/>'
            f'</w:num>'
        )

        numbering_elm.append(parse_xml(abstract_xml))
        numbering_elm.append(parse_xml(num_xml))
        self._bullet_num_id = next_num

    # ── Text helpers ─────────────────────────────────────────

    def _add_red_text(self, paragraph, text, bold=False, italic=False):
        """Add a run of red text to a paragraph."""
        run = paragraph.add_run(text)
        run.font.color.rgb = RED
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return run

    def _add_black_text(self, paragraph, text, bold=False, italic=False):
        """Add a run of black text to a paragraph."""
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return run

    def _set_cell_red(self, cell, text):
        """Set cell text to red (clearing existing content)."""
        p = cell.paragraphs[0]
        p.clear()
        self._add_red_text(p, text)

    def _set_cell_black(self, cell, text):
        """Set cell text to black (clearing existing content)."""
        p = cell.paragraphs[0]
        p.clear()
        self._add_black_text(p, text)

    def _set_cell_red_empty(self, cell):
        """Pre-format an empty cell as red so manual additions automatically appear red."""
        p = cell.paragraphs[0]
        p.clear()
        # Set default run properties on the paragraph to RED
        pPr = p._element.get_or_add_pPr()
        rPr_xml = (
            f'<w:rPr {nsdecls("w")}>'
            f'<w:color w:val="FF0000"/>'
            f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
            f'<w:sz w:val="22"/>'
            f'</w:rPr>'
        )
        existing = pPr.find(qn('w:rPr'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(parse_xml(rPr_xml))

    def _set_cell_shading(self, cell, fill_color):
        """Apply background shading to a cell."""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear" w:color="auto"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_cell_width(self, cell, width_dxa):
        """Set the width of a cell in dxa units."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
            tcPr.append(tcW)
        else:
            tcW.set(qn('w:w'), str(width_dxa))
            tcW.set(qn('w:type'), 'dxa')

    # ── Question text parsing ────────────────────────────────

    def _parse_question_text(self, question_text, selection_instruction):
        """Parse question text into formatted runs.

        Piped variables like <brand> are rendered in RED.
        Selection instruction is appended in italic BLACK.

        Returns list of (text, is_italic, is_red) tuples.
        """
        runs = []
        parts = re.split(r'(<[^>]+>)', question_text)
        for part in parts:
            if part.startswith('<') and part.endswith('>'):
                runs.append((part, False, True))   # piped variable → RED
            elif part:
                runs.append((part, False, False))  # normal text → BLACK
        runs.append((" ", False, False))
        if selection_instruction:
            runs.append((selection_instruction, True, False))  # italic BLACK
        return runs

    # ── Table helpers ────────────────────────────────────────

    def _set_table_borders(self, table):
        """Set BFBFBF borders on all sides of a table."""
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}/>'
        )
        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'</w:tblBorders>'
        )
        borders = parse_xml(borders_xml)
        existing = tbl_pr.find(qn('w:tblBorders'))
        if existing is not None:
            tbl_pr.remove(existing)
        tbl_pr.append(borders)

    def _fill_header_rows(self, table, num_cols, topic, q_number, text_runs, prog_note):
        """Fill the first 4 rows (R0-R3) of a table with merged header content.

        Cells in R0-R3 are horizontally merged across all columns so the header
        spans the full table width.
        """
        # Merge R0-R3 across all columns
        for r in range(4):
            if num_cols > 1:
                table.cell(r, 0).merge(table.cell(r, num_cols - 1))

        # R0: Topic — RED text on blue background
        r0 = table.cell(0, 0)
        self._set_cell_red(r0, topic)
        self._set_cell_shading(r0, BLUE_FILL)

        # R1: Question number — RED text
        r1 = table.cell(1, 0)
        if q_number:
            self._set_cell_red(r1, q_number)
        else:
            r1.text = ""

        # R2: Question text (piped variables RED) + selection instruction (italic BLACK)
        r2 = table.cell(2, 0)
        p = r2.paragraphs[0]
        p.clear()
        for text, is_italic, is_red in text_runs:
            if is_red:
                self._add_red_text(p, text, italic=is_italic)
            else:
                self._add_black_text(p, text, italic=is_italic)

        # R3: Programming note — RED text
        r3 = table.cell(3, 0)
        self._set_cell_red(r3, prog_note)

    def _build_header_table(self, topic, q_number, text_runs, prog_note):
        """Build a standalone 4-row x 1-col header table.

        Used for messages, dropdowns, and open-end questions that have no
        response table to merge into.
        """
        table = self.doc.add_table(rows=4, cols=1)
        self._set_table_borders(table)
        self._fill_header_rows(table, 1, topic, q_number, text_runs, prog_note)
        return table

    def _create_merged_table(self, total_rows, num_cols, topic, q_number, text_runs, prog_note):
        """Create a table with merged header rows (R0-R3) + response rows below."""
        table = self.doc.add_table(rows=total_rows, cols=num_cols)
        self._set_table_borders(table)
        self._fill_header_rows(table, num_cols, topic, q_number, text_runs, prog_note)
        return table

    def _fill_grid_header_cell(self, cell, header_text):
        """Fill a grid column header cell with split formatting.

        'C1:' prefix is RED, descriptive text after is BLACK.
        """
        p = cell.paragraphs[0]
        p.clear()
        match = re.match(r'(C\d+:\s*)(.*)', header_text)
        if match:
            self._add_red_text(p, match.group(1))
            if match.group(2):
                self._add_black_text(p, match.group(2))
        else:
            # No C-prefix — render entirely in red
            self._add_red_text(p, header_text)

    # ── NEW SCREEN ──────────────────────────────────────────

    def _add_new_screen(self):
        """Add 'NEW SCREEN' separator in red between question blocks."""
        self.doc.add_paragraph()  # blank line
        p = self.doc.add_paragraph()
        self._add_red_text(p, "NEW SCREEN")

    # ── Logic notes ─────────────────────────────────────────

    def _add_logic_note(self, text, bold=False):
        """Add a red logic/programming note paragraph after a question."""
        p = self.doc.add_paragraph()
        self._add_red_text(p, text, bold=bold)

    # ── List item helpers ───────────────────────────────────

    def _add_list_item(self, text, level=0, bold=False):
        """Add a properly formatted bullet list item at the given indentation level.

        Uses Word numbering definitions for real interactive bullets
        (level 0 = bullet, level 1 = sub-bullet, level 2 = sub-sub-bullet).
        """
        p = self.doc.add_paragraph(style='List Paragraph')

        # Set numPr for proper Word bullets
        pPr = p._element.get_or_add_pPr()
        numPr_xml = (
            f'<w:numPr {nsdecls("w")}>'
            f'<w:ilvl w:val="{level}"/>'
            f'<w:numId w:val="{self._bullet_num_id}"/>'
            f'</w:numPr>'
        )
        existing = pPr.find(qn('w:numPr'))
        if existing is not None:
            pPr.remove(existing)
        pPr.insert(0, parse_xml(numPr_xml))

        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        if bold:
            run.bold = True
        return p

    def _add_nested_list(self, items, level=0):
        """Add nested list items recursively.

        Each item is either:
          - str: plain bullet at current level
          - (str, [sub_items]): bullet with nested sub-items
        """
        for item in items:
            if isinstance(item, str):
                self._add_list_item(item, level=level)
            elif isinstance(item, tuple) and len(item) == 2:
                text, sub = item
                if isinstance(sub, list):
                    self._add_list_item(text, level=level)
                    self._add_nested_list(sub, level=level + 1)
                else:
                    # (text, note) — just a plain item with appended note
                    combined = f"{text} — {sub}" if sub else text
                    self._add_list_item(combined, level=level)

    # ── Study Info ──────────────────────────────────────────

    def set_study_info(self, title: str, date: str, doc_type: str = "Quantitative Survey Document"):
        """Add the title page header."""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        run = p.add_run(doc_type)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.color.rgb = BLACK

        p = self.doc.add_paragraph()
        run = p.add_run(date)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.color.rgb = BLACK

        self.doc.add_paragraph()  # blank line

    # ── Overview ────────────────────────────────────────────

    def add_overview(self, objectives: list, structure: list, criteria: list):
        """Add the Study Overview section with proper Word bullet formatting.

        Args:
            objectives: list of strings or (text, [sub_items]) tuples
            structure: list of (section_name, [sub_bullets]) tuples
            criteria: list of strings or (text, [sub_items]) tuples (supports nesting)
        """
        h = self.doc.add_heading("Survey Overview", level=1)
        for run in h.runs:
            run.font.size = Pt(12)

        # Research Objectives
        self.doc.add_heading("Research Objectives", level=2)
        self._add_nested_list(objectives)

        # Survey Structure
        self.doc.add_heading("Survey Structure", level=2)
        for section_name, sub_bullets in structure:
            self._add_list_item(section_name, level=0, bold=True)
            for sub in sub_bullets:
                self._add_list_item(sub, level=1)

        # Respondent Criteria & Quotas
        self.doc.add_heading("Respondent Criteria & Quotas", level=2)
        self._add_nested_list(criteria)

    # ── Section Headers ─────────────────────────────────────

    def add_section_header(self, section_name: str, objectives: list):
        """Add a section header with bulleted objectives."""
        self.doc.add_paragraph()
        self.doc.add_heading(section_name, level=2)

        p = self.doc.add_paragraph()
        run = p.add_run("Objectives:")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

        for obj in objectives:
            self._add_list_item(obj, level=0)

    # ── Hold Terminates Note ────────────────────────────────

    def add_hold_terminates_note(self):
        """Add the standard hold terminates note (RED, bold)."""
        p = self.doc.add_paragraph()
        self._add_red_text(
            p,
            "Please hold all terminations until end of screener unless otherwise specified.",
            bold=True
        )

    # ── Message Blocks ──────────────────────────────────────

    def add_message(self, topic: str, q_number: str, message_text: str, show_condition: str):
        """Add a message/transition screen (standalone 4x1 header table)."""
        if not self._first_block:
            self._add_new_screen()
        self._first_block = False

        text_runs = self._parse_question_text(message_text, "")
        self._build_header_table(
            topic=topic,
            q_number=q_number,
            text_runs=text_runs,
            prog_note=show_condition
        )

    # ── Question Blocks ─────────────────────────────────────

    def add_question(
        self,
        q_number: str,
        topic: str,
        question_text: str,
        selection_instruction: str,
        programming_note: str,
        response_options: list = None,
        grid_headers: list = None,
        logic_notes: list = None,
        question_type: str = "simple",
        scale_labels: list = None,
        bipolar_pairs: list = None,
    ):
        """Add a complete question block.

        For question types with response tables (simple, grid, scale, bipolar),
        the header rows (R0-R3) and response rows are merged into a single table.
        For dropdown/open_end, a standalone header table is used.

        Args:
            q_number: "S1.", "Q101.", "D1." — goes in R1 in RED
            topic: Topic label — goes in R0 with blue shading, RED text
            question_text: Question wording (BLACK, piped <variables> in RED)
            selection_instruction: e.g., "Select one." (italic BLACK)
            programming_note: Goes in R3 in RED
            response_options: List of tuples (format varies by question_type):
                - simple: [(option_text, note), ...]
                - 2col_grid: [(option_text, c1_note, c2_note), ...]
                - 3col_grid: [(option_text, c1_note, c2_note, c3_note), ...]
                - scale: [(statement_text, condition), ...] or [statement_text, ...]
            grid_headers: Column header strings for grids (e.g., "C1: Aware of")
            logic_notes: Logic instruction strings (RED paragraphs after table)
            question_type: "simple"|"2col_grid"|"3col_grid"|"scale"|"bipolar"|"dropdown"|"open_end"
            scale_labels: For "scale" type, list of scale point labels
            bipolar_pairs: For "bipolar" type, list of (left, right) tuples
        """
        if not self._first_block:
            self._add_new_screen()
        self._first_block = False

        text_runs = self._parse_question_text(question_text, selection_instruction)

        if question_type == "simple" and response_options:
            num_cols = 3
            total_rows = 4 + len(response_options)
            table = self._create_merged_table(
                total_rows, num_cols, topic, q_number, text_runs, programming_note
            )
            self._fill_simple_rows(table, response_options, start_row=4)

        elif question_type in ("2col_grid", "3col_grid") and response_options:
            num_data_cols = 2 if question_type == "2col_grid" else 3
            num_cols = 2 + num_data_cols
            # +1 for column header row (C1, C2, etc.)
            total_rows = 4 + 1 + len(response_options)
            table = self._create_merged_table(
                total_rows, num_cols, topic, q_number, text_runs, programming_note
            )
            self._fill_grid_rows(
                table, response_options, grid_headers, num_data_cols, header_row=4
            )

        elif question_type == "scale" and scale_labels:
            # Merged header + scale numbers + scale labels (one table)
            num_points = len(scale_labels)
            total_rows = 4 + 2  # R0-R3 header + numbers row + labels row
            table = self._create_merged_table(
                total_rows, num_points, topic, q_number, text_runs, programming_note
            )
            self._fill_scale_header(table, scale_labels, start_row=4)
            # Separate statement table below (if statements provided)
            if response_options:
                self._add_scale_statements(response_options)

        elif question_type == "bipolar" and bipolar_pairs:
            num_cols = 3  # blank | C1 | C2
            total_rows = 4 + 1 + len(bipolar_pairs)  # header + C1/C2 row + pairs
            table = self._create_merged_table(
                total_rows, num_cols, topic, q_number, text_runs, programming_note
            )
            self._fill_bipolar_rows(table, bipolar_pairs, header_row=4)

        else:
            # dropdown, open_end — standalone header table only
            self._build_header_table(topic, q_number, text_runs, programming_note)

        # Logic notes (RED paragraphs after the table)
        if logic_notes:
            for note in logic_notes:
                self._add_logic_note(note)

    # ── Response Row Builders ───────────────────────────────

    def _fill_simple_rows(self, table, options, start_row):
        """Fill response rows in a simple 3-col merged table.

        Col 0: blank (red-formatted for manual use)
        Col 1: option text (BLACK)
        Col 2: note (RED, pre-formatted even if empty)
        """
        for i, (opt_text, note) in enumerate(options):
            row = start_row + i

            # Col 0: blank, red-formatted
            c0 = table.cell(row, 0)
            self._set_cell_red_empty(c0)
            self._set_cell_width(c0, COL0_WIDTH)

            # Col 1: option text (BLACK)
            c1 = table.cell(row, 1)
            self._set_cell_black(c1, opt_text)
            self._set_cell_width(c1, COL1_WIDTH)

            # Col 2: note (RED) — always red-formatted even if empty
            c2 = table.cell(row, 2)
            if note:
                self._set_cell_red(c2, note)
            else:
                self._set_cell_red_empty(c2)
            self._set_cell_width(c2, COL2_WIDTH)

    def _fill_grid_rows(self, table, options, headers, num_data_cols, header_row):
        """Fill grid column header row + data rows.

        header_row: row index for column headers (C1, C2, etc.)
        Data rows start at header_row + 1.

        Column headers: "C1:" prefix in RED, descriptive text in BLACK.
        """
        total_cols = 2 + num_data_cols

        # Column header row
        table.cell(header_row, 0).text = ""
        table.cell(header_row, 1).text = ""
        if headers:
            for h_idx, h_text in enumerate(headers):
                col_idx = h_idx + 2
                if col_idx < total_cols:
                    self._fill_grid_header_cell(table.cell(header_row, col_idx), h_text)

        # Data rows
        data_start = header_row + 1
        for i, option_data in enumerate(options):
            row = data_start + i

            # Col 0: blank, red-formatted
            self._set_cell_red_empty(table.cell(row, 0))
            self._set_cell_width(table.cell(row, 0), COL0_WIDTH)

            # Col 1: option text (BLACK)
            self._set_cell_black(table.cell(row, 1), option_data[0])

            # Remaining columns: notes (RED, pre-formatted even if empty)
            for c in range(1, len(option_data)):
                col_idx = c + 1
                if col_idx < total_cols:
                    if option_data[c]:
                        self._set_cell_red(table.cell(row, col_idx), option_data[c])
                    else:
                        self._set_cell_red_empty(table.cell(row, col_idx))

    def _fill_scale_header(self, table, scale_labels, start_row):
        """Fill scale number row and label row within a merged table.

        start_row: row for numbers (RED)
        start_row + 1: row for labels (BLACK, label text only — no number prefix)
        """
        for i, label in enumerate(scale_labels):
            # Numbers row: RED
            self._set_cell_red(table.cell(start_row, i), str(i + 1))
            # Labels row: BLACK (label text only, no number prefix)
            self._set_cell_black(table.cell(start_row + 1, i), label)

    def _add_scale_statements(self, statements):
        """Add a separate statement table for scale batteries.

        Format: N rows x 2 cols (blank | statement).
        If any statement has a condition, adds a 3rd column (RED).

        Args:
            statements: list of strings, or list of (text, condition) tuples
        """
        # Normalize to tuples
        normalized = []
        for stmt in statements:
            if isinstance(stmt, str):
                normalized.append((stmt, ""))
            else:
                normalized.append((stmt[0], stmt[1] if len(stmt) > 1 else ""))

        has_conditions = any(cond for _, cond in normalized)
        num_cols = 3 if has_conditions else 2

        stmt_table = self.doc.add_table(rows=len(normalized), cols=num_cols)
        self._set_table_borders(stmt_table)

        for i, (stmt_text, condition) in enumerate(normalized):
            # Col 0: blank, red-formatted
            self._set_cell_red_empty(stmt_table.cell(i, 0))
            self._set_cell_width(stmt_table.cell(i, 0), COL0_WIDTH)

            # Col 1: statement text (BLACK)
            self._set_cell_black(stmt_table.cell(i, 1), stmt_text)

            # Col 2: condition (RED) — only if column exists
            if has_conditions:
                if condition:
                    self._set_cell_red(stmt_table.cell(i, 2), condition)
                else:
                    self._set_cell_red_empty(stmt_table.cell(i, 2))

    def _fill_bipolar_rows(self, table, pairs, header_row):
        """Fill bipolar/slider scale rows.

        header_row: row for C1/C2 column labels (RED)
        Data rows: blank | left statement (BLACK) | right statement (BLACK)

        The programmer builds the actual numeric scale from the programming note.
        """
        # C1/C2 header row
        self._set_cell_red_empty(table.cell(header_row, 0))
        p = table.cell(header_row, 1).paragraphs[0]
        p.clear()
        self._add_red_text(p, "C1")
        p = table.cell(header_row, 2).paragraphs[0]
        p.clear()
        self._add_red_text(p, "C2")

        # Data rows: left statement | right statement
        data_start = header_row + 1
        for i, (left, right) in enumerate(pairs):
            row = data_start + i
            self._set_cell_red_empty(table.cell(row, 0))
            self._set_cell_width(table.cell(row, 0), COL0_WIDTH)
            self._set_cell_black(table.cell(row, 1), left)
            self._set_cell_black(table.cell(row, 2), right)

    # ── Save ────────────────────────────────────────────────

    def save(self, filepath: str):
        """Save the document to the specified path."""
        directory = os.path.dirname(filepath)
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

        self.doc.save(filepath)
        print(f"Survey document saved to: {os.path.abspath(filepath)}")
        return filepath


# ── Wireframe Parser ────────────────────────────────────────

def parse_wireframe(filepath: str) -> dict:
    """
    Parse a wireframe .docx and extract structured question data.

    Returns:
        dict with keys:
            - title: str
            - overview: dict with objectives, structure, criteria
            - sections: list of dicts, each with name, objectives, questions
    """
    doc = Document(filepath)

    result = {
        'title': '',
        'overview': {'objectives': [], 'structure': [], 'criteria': []},
        'sections': []
    }

    for p in doc.paragraphs[:5]:
        text = p.text.strip()
        if text and not result['title']:
            result['title'] = text
            break

    for table in doc.tables:
        section = _parse_wireframe_table(table)
        if section and section['questions']:
            result['sections'].append(section)

    return result


def _parse_wireframe_table(table) -> dict:
    """Parse a single wireframe table into a section dict."""
    rows = table.rows
    if len(rows) < 2:
        return None

    header_cells = [cell.text.strip().lower() for cell in rows[0].cells]
    num_cols = len(header_cells)

    topic_col = question_col = response_col = objective_col = number_col = None

    for i, h in enumerate(header_cells):
        if 'topic' in h:
            topic_col = i
        elif 'question' in h and 'number' not in h:
            question_col = i
        elif 'response' in h or 'sample' in h:
            response_col = i
        elif 'objective' in h or 'rationale' in h:
            objective_col = i
        elif 'number' in h or '#' in h:
            number_col = i

    if topic_col is None and num_cols >= 3:
        if number_col is not None:
            topic_col = topic_col or 1
            question_col = question_col or 2
            response_col = response_col or 3
            objective_col = objective_col or (4 if num_cols > 4 else None)
        else:
            topic_col = 0
            question_col = 1
            response_col = 2
            objective_col = 3 if num_cols > 3 else None

    if topic_col is None or question_col is None:
        return None

    section_name = ""
    questions = []

    for row_idx in range(1, len(rows)):
        cells = [cell.text.strip() for cell in rows[row_idx].cells]

        topic = cells[topic_col] if topic_col < len(cells) else ""
        question_text = cells[question_col] if question_col is not None and question_col < len(cells) else ""
        response = cells[response_col] if response_col is not None and response_col < len(cells) else ""
        objective = cells[objective_col] if objective_col is not None and objective_col < len(cells) else ""
        number = cells[number_col] if number_col is not None and number_col < len(cells) else ""

        if not topic and not question_text:
            continue

        if topic and not question_text and not response:
            if not section_name:
                section_name = topic
            continue

        questions.append({
            'number': number,
            'topic': topic,
            'question_text': question_text,
            'response_options_raw': response,
            'objective_rationale': objective,
        })

    return {
        'name': section_name or "Section",
        'objectives': [],
        'questions': questions,
    }


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python survey_doc_builder.py <wireframe.docx> [output.docx]")
        sys.exit(1)

    wireframe_path = sys.argv[1]
    result = parse_wireframe(wireframe_path)

    print(f"Title: {result['title']}")
    print(f"Sections: {len(result['sections'])}")
    for sec in result['sections']:
        print(f"\n  {sec['name']} ({len(sec['questions'])} questions)")
        for q in sec['questions']:
            topic = q['topic'][:40]
            qtext = q['question_text'][:60]
            print(f"    {q['number']:>4} | {topic:<40} | {qtext}")
